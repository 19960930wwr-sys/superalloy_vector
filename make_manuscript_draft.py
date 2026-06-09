"""Generate a Word manuscript draft for the superalloy multi-task project."""
from pathlib import Path
import json

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "output" / "draft_materials"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "superalloy_multitask_inverse_design_manuscript_draft.docx"

TASK_LABELS = {
    "density": "Density",
    "creep": "Creep life",
    "liquidus": "Liquidus temperature",
    "phase_class": "Harmful phase probability",
    "size": "Gamma-prime size",
    "solidus": "Solidus temperature",
    "solvus": "Gamma-prime solvus temperature",
}


def fmt(value, nd=3):
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    if isinstance(value, (int, np.integer)):
        return str(int(value))
    if isinstance(value, (float, np.floating)):
        if abs(value) >= 100:
            return f"{value:.1f}"
        return f"{value:.{nd}f}"
    return str(value)


def set_cell_text(cell, text, bold=False, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, df, caption=None, font_size=8.2, header_fill="D9EAF7"):
    if caption:
        p = doc.add_paragraph()
        p.style = doc.styles["Caption"]
        r = p.add_run(caption)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.5)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        set_cell_text(cell, col, bold=True, size=font_size)
        shade_cell(cell, header_fill)
    for i in range(df.shape[0]):
        for j, col in enumerate(df.columns):
            set_cell_text(table.cell(i + 1, j), fmt(df.iloc[i, j]), size=font_size)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
    return p


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.8)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)


def add_figure(doc, rel_path, caption, width_cm=15.0):
    fig = ROOT / rel_path
    if not fig.exists():
        add_paragraph(doc, f"[Figure file missing: {rel_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(fig), width=Cm(width_cm))
    cp = doc.add_paragraph()
    cp.style = doc.styles["Caption"]
    cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = cp.add_run(caption)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.5)


def normalize_styles(doc):
    for style_name in ["Normal", "Caption"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for style_name, size in [
        ("Title", 18),
        ("Heading 1", 14),
        ("Heading 2", 12.5),
        ("Heading 3", 11.5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def load_data():
    mt = pd.read_csv(ROOT / "data" / "master_table.csv")
    comp = pd.read_csv(ROOT / "output" / "results" / "full_comparison_table.csv")
    best = pd.read_csv(ROOT / "output" / "results" / "best_model_per_task_embedding.csv")
    top = pd.read_csv(
        ROOT
        / "output"
        / "results"
        / "forward_screen_seedtop10_seeds_T1080-1120_S120-160_life270.csv"
    )
    corpus_stats = json.loads(
        (ROOT / "output" / "corpus_processed" / "corpus_stats.json").read_text(
            encoding="utf-8"
        )
    )
    with open(ROOT / "output" / "results" / "task_grouping_E_pa.json", encoding="utf-8") as f:
        grp_pa = json.load(f)
    with open(ROOT / "output" / "results" / "task_grouping_E_base.json", encoding="utf-8") as f:
        grp_base = json.load(f)
    return mt, comp, best, top, corpus_stats, grp_pa, grp_base


def make_tables(mt, comp, best, top, corpus_stats, grp_pa, grp_base):
    stats = mt.groupby("task")["target"].agg(["count", "mean", "std", "min", "max"]).reset_index()
    stats["Property"] = stats["task"].map(TASK_LABELS)
    task_types = mt.drop_duplicates("task").set_index("task")["task_type"].to_dict()
    stats["Task type"] = stats["task"].map(task_types)
    stats_table = stats[
        ["task", "Property", "Task type", "count", "mean", "std", "min", "max"]
    ].rename(
        columns={
            "task": "Task",
            "count": "N",
            "mean": "Mean",
            "std": "Std",
            "min": "Min",
            "max": "Max",
        }
    )

    full = comp[(comp["model"] == "multitask") & (comp["embedding"].isin(["E_pa", "E_base"]))].copy()
    rows = []
    for task in ["density", "creep", "liquidus", "phase_class", "size", "solidus", "solvus"]:
        row = {"Task": task}
        for emb in ["E_pa", "E_base"]:
            sub = full[(full["task"] == task) & (full["embedding"] == emb)].iloc[0]
            if sub["task_type"] == "classification":
                row[f"{emb} metric"] = "F1"
                row[f"{emb} value"] = sub["f1_mean"]
                row[f"{emb} std"] = sub["f1_std"]
            else:
                row[f"{emb} metric"] = "R2"
                row[f"{emb} value"] = sub["r2_mean"]
                row[f"{emb} std"] = sub["r2_std"]
        row["Delta"] = row["E_pa value"] - row["E_base value"]
        rows.append(row)
    full_table = pd.DataFrame(rows)[
        ["Task", "E_pa metric", "E_pa value", "E_pa std", "E_base value", "E_base std", "Delta"]
    ]

    grouped = comp[
        comp["model"].isin(["grouped-paG1", "grouped-paG2", "grouped-baseG1", "grouped-baseG2"])
    ].copy()
    rows = []
    for _, r in grouped.iterrows():
        metric = "F1" if r["task_type"] == "classification" else "R2"
        val = r["f1_mean"] if metric == "F1" else r["r2_mean"]
        mtv_sub = full[(full["task"] == r["task"]) & (full["embedding"] == r["embedding"])].iloc[0]
        mtv = mtv_sub["f1_mean"] if metric == "F1" else mtv_sub["r2_mean"]
        rows.append(
            {
                "Embedding": r["embedding"],
                "Grouped model": r["model"],
                "Task": r["task"],
                "Metric": metric,
                "Grouped": val,
                "Full MT": mtv,
                "Delta": val - mtv,
            }
        )
    grouped_table = pd.DataFrame(rows).sort_values(["Embedding", "Grouped model", "Task"])

    best_table = best.copy()
    best_table["Task"] = best_table["task"]
    best_table["Metric"] = best_table["metric"]
    best_table = best_table[
        ["Task", "embedding", "best_model", "Metric", "best_value", "multitask_value", "delta_vs_multitask"]
    ].rename(
        columns={
            "embedding": "Embedding",
            "best_model": "Best model",
            "best_value": "Best value",
            "multitask_value": "Full MT",
            "delta_vs_multitask": "Delta",
        }
    )

    props = top[
        [
            "creep_real",
            "test_temp",
            "test_stress",
            "solvus",
            "processing_window",
            "density",
            "phase_class",
            "size",
            "freezing_range",
            "score",
        ]
    ].copy()
    props.insert(0, "Rank", np.arange(1, len(props) + 1))
    props = props.rename(
        columns={
            "creep_real": "Creep h",
            "test_temp": "Test degC",
            "test_stress": "Stress MPa",
            "solvus": "Solvus degC",
            "processing_window": "Window degC",
            "density": "Density",
            "phase_class": "Phase prob",
            "size": "Size",
            "freezing_range": "Freezing degC",
            "score": "Score",
        }
    )

    main_elems = [c for c in ["Ni", "Co", "Al", "W", "Ta", "Mo", "Re", "Cr", "Ru", "Hf", "Ti", "Nb"] if c in top.columns]
    comps_top = top[main_elems].copy()
    comps_top.insert(0, "Rank", np.arange(1, len(comps_top) + 1))

    pretrain_table = pd.DataFrame(
        [
            ["Documents in corpus", corpus_stats["num_documents"]],
            ["Sentences after splitting", corpus_stats["num_sentences"]],
            ["Valid tokenized sentences", corpus_stats["num_valid_tokenized"]],
            ["Vocabulary size", corpus_stats["vocab_size"]],
            ["Average tokens per sentence", round(corpus_stats["avg_tokens_per_sentence"], 2)],
            ["Embedding dimension", 128],
            ["BERT-base MLM final loss", 2.0406],
            ["PA-MLM final total loss", 2.6431],
            ["PA-MLM final MLM loss", 2.1770],
            ["PA-MLM final element-attribute loss", 0.3900],
            ["PA-MLM final process-classification loss", 0.1521],
        ],
        columns=["Item", "Value"],
    )

    group_summary = pd.DataFrame(
        [
            [
                "E_pa",
                "M3",
                grp_pa["M3"]["best_k"],
                "; ".join([f"{g}: {', '.join(v)}" for g, v in grp_pa["final_groups"].items()]),
            ],
            [
                "E_base",
                "M3",
                grp_base["M3"]["best_k"],
                "; ".join([f"{g}: {', '.join(v)}" for g, v in grp_base["final_groups"].items()]),
            ],
        ],
        columns=["Embedding", "Evidence", "Best k", "Final task groups"],
    )

    cascade_table = pd.DataFrame(
        [
            ["Initial high-creep seeds near 1080-1120 degC and 120-160 MPa", 52],
            ["Solvus >= 1220 degC", 52],
            ["Density <= 8.9 g/cm3", 52],
            ["Processing window >= 80 degC", 18],
            ["Phase probability <= 0.5", 11],
            ["Gamma-prime size <= 500", 11],
            ["Freezing range <= 60 degC", 11],
            ["Top candidates selected by weighted Z-score", 10],
        ],
        columns=["Filtering step", "Remaining candidates"],
    )
    return {
        "stats": stats_table,
        "pretrain": pretrain_table,
        "group_summary": group_summary,
        "full": full_table,
        "grouped": grouped_table,
        "best": best_table,
        "cascade": cascade_table,
        "props": props,
        "comps_top": comps_top,
    }


def build_doc(tables, top):
    doc = Document()
    normalize_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Semantic property-aware embeddings and multi-task learning enable inverse design of high-temperature superalloys"
    )
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(17)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Draft manuscript prepared from the local superalloy-vector project files")
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Author list and affiliations: to be completed")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)

    add_heading(doc, "Abstract", 1)
    add_paragraph(
        doc,
        "Data-driven design of high-temperature superalloys requires models that can reason over composition, heat treatment, creep test conditions, phase stability, transition temperatures, density, and microstructural length scales within a single framework. In this work, a domain-specific semantic representation was developed from a corpus of 26,980 superalloy full-text documents containing 4,753,513 valid tokenized sentences and a vocabulary of 74,258 tokens. A property-aware masked language model (PA-MLM) was trained by combining conventional masked-token reconstruction with auxiliary elemental-property regression and process-category prediction. The resulting embedding, denoted E_pa, was integrated with numerical composition and processing descriptors to train a seven-task superalloy model for density, creep life, liquidus temperature, solidus temperature, gamma-prime solvus temperature, gamma-prime size, and harmful phase classification.",
    )
    add_paragraph(
        doc,
        "A master table containing 3,911 records was assembled from seven task-specific datasets. Under five-fold cross-validation, the full multi-task neural network using E_pa outperformed the corresponding plain BERT-MLM embedding (E_base) on all seven downstream tasks. The E_pa model achieved R2 values of 0.806 for density, 0.736 for creep life, 0.895 for liquidus temperature, 0.946 for gamma-prime size, 0.853 for solidus temperature, and 0.875 for gamma-prime solvus temperature, together with an F1 score of 0.930 for phase classification. Task-affinity analysis using Ridge-weight similarity and shared-gradient similarity further showed that E_pa flattens inter-task gradient conflicts, whereas E_base depends strongly on a temperature-task cluster. Finally, a seed-based inverse-design protocol combined experimentally measured high-creep-life alloys with model-predicted thermophysical constraints. From 52 high-creep seeds, sequential constraints on solvus, density, processing window, phase probability, gamma-prime size, and freezing range yielded 11 feasible candidates, from which the top 10 were recommended. The leading alloy has a measured creep life of 722 h at 1100 degC/137 MPa and predicted solvus of 1255.8 degC, processing window of 93.9 degC, density of 8.78 g/cm3, harmful phase probability of 0.41, gamma-prime size of 340, and freezing range of 31.9 degC.",
    )
    add_paragraph(
        doc,
        "Keywords: superalloys; multi-task learning; masked language model; property-aware embedding; creep life; inverse design; seed-based screening",
    )

    add_heading(doc, "1. Introduction", 1)
    add_paragraph(
        doc,
        "High-temperature superalloys are used in turbine blades, vanes, combustor components, and other severe-service environments because they combine high-temperature strength, creep resistance, phase stability, and oxidation resistance. Their exceptional performance originates from controlled combinations of matrix chemistry, gamma-prime precipitation, refractory additions, heat treatment, and microstructural stability. The same ingredients that make these alloys powerful also make their design difficult. A composition that raises the gamma-prime solvus may also increase density, narrow the heat-treatment window, promote deleterious phases, or alter the freezing range. A practical alloy design therefore cannot optimize a single response in isolation; it must balance coupled objectives across thermodynamics, processing, microstructure, and mechanical performance.",
    )
    add_paragraph(
        doc,
        "Machine learning offers an appealing route for accelerating such searches, but several obstacles remain. Superalloy data are heterogeneous: density and liquidus measurements may not contain explicit process information, creep data depend strongly on temperature and stress, and gamma-prime size is meaningful only with a specified heat-treatment path. Single-task models neglect useful physical coupling among related properties, whereas naive multi-task learning may suffer negative transfer when unrelated tasks compete for the same shared representation. Conventional tabular descriptors also encode elements as independent columns and do not capture the semantic and physical proximity learned from scientific literature.",
    )
    add_paragraph(
        doc,
        "The present work addresses these problems through a four-stage pipeline. First, a large superalloy literature corpus was processed with a domain tokenizer that preserves element symbols and process phrases. Second, two language-model embeddings were trained: a plain BERT-MLM embedding (E_base) and a property-aware embedding (E_pa) that explicitly learns elemental physical attributes and process categories during masked-language pretraining. Third, these semantic embeddings were fused with numerical composition and process information in a multi-task learning framework covering seven superalloy property tasks. Fourth, the trained ensemble was used in a conservative seed-based inverse design strategy, in which measured high-creep-life alloys were retained as reliable anchors and the model predicted the remaining thermophysical and microstructural constraints.",
    )

    add_heading(doc, "2. Methodology", 1)
    add_heading(doc, "2.1. Literature corpus and task datasets", 2)
    add_paragraph(
        doc,
        "The language-model pretraining corpus consisted of full-text superalloy articles stored as plain-text files. After document reading, sentence splitting, and domain tokenization, the final processed corpus contained 26,980 documents, 4,754,807 sentences, and 4,753,513 valid tokenized samples. The average sentence length was 27.21 tokens, and the vocabulary contained 74,258 unique tokens.",
    )
    add_paragraph(
        doc,
        "The BERT pretraining literature was constructed as a domain corpus rather than a generic materials-text collection. Full-text articles were programmatically archived from CrossRef, Web of Science, Scopus, and ScienceDirect retrieval channels. For HTML articles, ChemDataExtractor-based scraping was used to recover the article body. During preprocessing, only abstracts and main-body text were retained, while DOI records, article identifiers, titles, journal metadata, publication information, URLs, and other boilerplate text were removed. A document-level keyword audit of the raw text archive before final validity filtering (27,256 files) confirmed that the corpus was dominated by superalloy and high-temperature-alloy literature: 22,766 files contained superalloy terminology, 18,394 contained high-temperature expressions, 16,417 contained Ni-based or nickel-based descriptions, 12,086 contained creep or rupture terms, and 10,897 contained solution-treatment or aging vocabulary. The corpus also included documents related to Co-based and Co-Ni superalloys, gamma-prime strengthening, solvus and solidus/liquidus temperatures, TCP or other harmful phases, refractory alloying, solidification, heat treatment, and microstructure-property relationships. This literature scope allowed the BERT model to learn metallurgical co-occurrence patterns linking element symbols, process verbs, test conditions, units, phase names, and property labels before being coupled with numerical descriptors for the downstream tasks.",
    )
    add_paragraph(
        doc,
        "The downstream supervised dataset was assembled from seven task-specific spreadsheets: density, creep life, liquidus temperature, solidus temperature, gamma-prime solvus temperature, gamma-prime size, and harmful phase classification. The individual files were standardized into a long-format master table. Each row contains a task name, a 24-element composition vector, six possible process or test-condition variables, a target value, a target name, and the task type.",
    )
    add_table(doc, tables["stats"], "Table 1. Downstream task statistics in the unified master table.", 7.8)

    add_heading(doc, "2.2. Domain tokenizer and property-aware masked language model", 2)
    add_paragraph(
        doc,
        "A superalloy-specific tokenizer was implemented to avoid common tokenization errors in materials text. Element symbols are preserved as case-sensitive atomic tokens, so that Ni, Co, Al, C, and B are not confused with ordinary words or subwords. Multi-word process phrases, such as solution treatment, aging, creep, rupture, homogenization, hot isostatic pressing, directional solidification, and single crystal, are protected before ordinary tokenization. Numerical values are mapped to a generic numerical token to reduce sparsity while preserving sentence context.",
    )
    add_paragraph(
        doc,
        "Two transformer embeddings were trained. The E_base model used a compact BERT-style masked language model with hidden size 128, six transformer layers, eight attention heads, intermediate size 512, maximum sequence length 128, dropout 0.1, and mixed-precision training. The PA-MLM model used the same backbone, but added two auxiliary heads. The first predicts a normalized 13-dimensional physicochemical attribute vector for masked element tokens. The second predicts the process category for masked process tokens. The total PA-MLM loss was L = L_MLM + lambda_attr L_attr + lambda_proc L_proc, where lambda_attr = 1.0 and lambda_proc = 0.5.",
    )
    add_table(doc, tables["pretrain"], "Table 2. Corpus statistics and pretraining summary.", 8.0)

    add_heading(doc, "2.3. Alloy representation", 2)
    add_paragraph(
        doc,
        "For each alloy, the composition representation was computed as a weighted average of elemental token embeddings. If w_i is the atomic percentage of element i and e_i is the corresponding learned token embedding, the composition vector is c = sum_i w_i e_i / sum_i w_i. Process information was encoded by combining an averaged process-action embedding with normalized numerical process parameters. The final input vector has dimension 263: 128 for composition, 128 for process action, six for numerical process parameters, and one process mask.",
    )
    add_paragraph(
        doc,
        "Regression targets were normalized by task-specific z-scores during neural-network training and denormalized for evaluation. The classification target for harmful phase formation was kept as a binary label and modeled with a sigmoid output. This task-wise normalization step was essential because the target scales differ by several orders of magnitude, from density around 8-10 g/cm3 to creep life that extends from less than 1 h to more than 6000 h.",
    )

    add_heading(doc, "2.4. Multi-task model and baselines", 2)
    add_paragraph(
        doc,
        "The full multi-task model contains a shared multilayer perceptron backbone followed by task-specific prediction heads. The backbone dimensions are 256, 128, and 64 with batch normalization, ReLU activation, and dropout of 0.2. Each regression head outputs one scalar in normalized target space; the classification head outputs a harmful-phase probability. Training used AdamW, learning rate 1e-3, batch size 128, five-fold cross-validation, and early stopping. SmoothL1 loss was used for regression, and binary cross entropy was used for phase classification. Homoscedastic uncertainty weighting was used to balance task losses, and the classification loss was amplified to prevent the single classification task from being overwhelmed by the six regression tasks.",
    )
    add_paragraph(
        doc,
        "Three model families were benchmarked against the full multi-task network: single-task neural networks, grouped multi-task networks, and classical machine-learning baselines. The classical baselines included Ridge regression, k-nearest-neighbor regression, support vector regression, random forest regression, gradient boosting regression, and corresponding classification variants.",
    )

    add_heading(doc, "2.5. Data-driven task grouping", 2)
    add_paragraph(
        doc,
        "Task grouping was performed using two independent evidence sources. The first method, M2, fits a simple linear model to each task and compares the cosine similarity between weight vectors. The second method, M3, uses a trained full multi-task model and records the gradient vector of each task with respect to shared backbone parameters. The cosine similarity between task gradients measures whether two tasks update the shared representation in compatible directions. Ward-linkage hierarchical clustering was then applied to the 1 - similarity distance matrix, and the number of groups was chosen by silhouette score over k = 2, 3, and 4.",
    )
    add_table(doc, tables["group_summary"], "Table 3. Data-driven task grouping derived from shared-gradient similarity (M3).", 8.0)

    add_heading(doc, "2.6. Seed-based inverse design and screening", 2)
    add_paragraph(
        doc,
        "Several inverse-design routes were implemented, including NSGA-II optimization and grid-based forward screening in predefined composition spaces. These routes are useful for exploring the response surface, but fully model-driven search can become unreliable in sparse regions of the training distribution, particularly for creep life under severe test conditions. Therefore, the final recommendation strategy used a conservative seed-based screening protocol. Measured creep life was treated as the reliability anchor, while the multi-task ensemble predicted the remaining properties needed for engineering feasibility.",
    )
    add_paragraph(
        doc,
        "The final seed condition was test temperature in the range 1080-1120 degC, stress in the range 120-160 MPa, and measured creep life above 270 h. The five-fold E_pa multi-task ensemble predicted solvus, density, harmful phase probability, liquidus, solidus, and gamma-prime size for each seed. Liquidus and solidus were converted into freezing range, and solidus minus solvus was converted into processing window. Hard constraints were applied sequentially: solvus >= 1220 degC, density <= 8.9 g/cm3, processing window >= 80 degC, harmful phase probability <= 0.5, gamma-prime size <= 500, and freezing range <= 60 degC.",
    )
    add_table(doc, tables["cascade"], "Table 4. Sequential filtering cascade for the final seed-based inverse-design recommendation.", 8.3)

    add_heading(doc, "3. Results and discussion", 1)
    add_heading(doc, "3.1. Pretraining behavior and representation construction", 2)
    add_paragraph(
        doc,
        "Both E_base and E_pa were trained for 15 epochs on 4.75 million tokenized sentences. The plain BERT-MLM model contained approximately 10.80 million parameters and reduced the average masked-language loss from 4.96 in the first epoch to 2.04 in the final epoch. The PA-MLM model contained approximately 10.84 million parameters and reduced the total loss from 6.34 to 2.64. At the end of training, the PA-MLM loss decomposed into MLM loss of 2.18, elemental attribute loss of 0.39, and process-category loss of 0.15.",
    )
    add_figure(
        doc,
        "output/figures/element_clustering.png",
        "Figure 1. Element clustering based on the learned semantic embedding. The plot provides a qualitative check that chemically and metallurgically related elements are organized in the learned representation space.",
        13.5,
    )

    add_heading(doc, "3.2. Full multi-task prediction performance", 2)
    add_paragraph(
        doc,
        "The most direct comparison is between full multi-task models trained with E_pa and E_base. E_pa achieved a strict win on all seven tasks. For creep life, E_pa increased R2 from 0.653 to 0.736 and reduced RMSE from 429.5 h to 369.6 h. For solvus, E_pa increased R2 from 0.846 to 0.875 and reduced RMSE from 35.6 degC to 32.0 degC. For phase classification, E_pa increased F1 from 0.906 to 0.930 while maintaining a high AUC near 0.979.",
    )
    add_table(doc, tables["full"], "Table 5. Full multi-task performance comparison between E_pa and E_base. Regression tasks report R2; phase classification reports F1.", 7.8)
    add_figure(doc, "output/figures/comparison_regression_r2.png", "Figure 2. Regression-task R2 comparison for full multi-task and single-task neural models using E_pa and E_base.", 14.0)
    add_figure(doc, "output/figures/comparison_classification.png", "Figure 3. Classification performance comparison for the harmful phase task.", 13.5)

    add_heading(doc, "3.3. Task similarity and grouped multi-task learning", 2)
    add_paragraph(
        doc,
        "The M3 task-gradient matrices revealed qualitatively different sharing behavior for E_pa and E_base. For E_pa, off-diagonal gradient cosine values were comparatively flat, ranging from -0.19 to 0.38. This indicates that property-aware pretraining reduced strong task conflicts and reduced the need for highly specialized task families. In contrast, E_base showed a strong temperature-task cluster: liquidus-solidus similarity reached 0.93, solidus-solvus reached 0.88, and liquidus-solvus reached 0.78.",
    )
    add_figure(doc, "output/figures/fig_heatmap_M3_compare.png", "Figure 4. Shared-gradient similarity (M3) heatmaps comparing E_pa and E_base.", 15.5)
    add_paragraph(
        doc,
        "Grouped multi-task learning was helpful only under specific conditions. In E_base, the liquidus-solidus-solvus group improved liquidus by 0.038 R2 and solidus by 0.050 R2 relative to the full multi-task baseline, confirming the M3 prediction that these temperature tasks share compatible gradients. In E_pa, the density-solidus-solvus group improved density from 0.806 to 0.841 and also slightly improved solidus and solvus. However, the E_pa group containing creep, liquidus, phase classification, and size degraded all four tasks.",
    )
    add_table(doc, tables["grouped"], "Table 6. Grouped multi-task models compared with their full multi-task baselines.", 7.5)
    add_figure(doc, "output/figures/fig_grouped_vs_full_bar.png", "Figure 5. Per-task comparison among grouped multi-task, full multi-task, single-task, and best classical machine-learning baselines.", 15.5)

    add_heading(doc, "3.4. Best model per task and remaining difficult targets", 2)
    add_paragraph(
        doc,
        "The best model family depended on the task. Under E_pa, full multi-task learning was already optimal for liquidus, grouped multi-task learning was optimal for density, and single-task learning was optimal for creep, size, solidus, and solvus. The best E_pa values were 0.841 R2 for density, 0.757 for creep life, 0.895 for liquidus, 0.935 F1 for phase classification, 0.953 for size, 0.883 for solidus, and 0.901 for solvus. Creep remained the most difficult target, which is physically reasonable because creep life depends on composition, temperature, stress, microstructure, and long-tailed rupture statistics.",
    )
    add_table(doc, tables["best"], "Table 7. Best model per task and embedding.", 7.6)
    add_figure(doc, "output/figures/fig_radar_best_per_task.png", "Figure 6. Radar plot of best per-task performance for E_pa and E_base.", 14.5)
    add_figure(doc, "output/figures/fig_delta_heatmap.png", "Figure 7. Performance change relative to the full multi-task baseline.", 15.0)

    add_heading(doc, "3.5. Seed-based inverse design results", 2)
    add_paragraph(
        doc,
        "The final candidate recommendation was obtained from measured high-creep seeds rather than a purely extrapolative optimizer. The initial seed set contained 52 alloys satisfying the target creep window near 1080-1120 degC and 120-160 MPa with measured life above 270 h. All 52 passed the solvus and density constraints. The processing-window criterion was the bottleneck: when the window threshold was set to 80 degC, 18 candidates remained. The harmful phase constraint reduced the set to 11 candidates, and the size and freezing-range constraints retained the same 11. The top 10 were then selected by weighted multi-objective ranking.",
    )
    add_table(doc, tables["props"], "Table 8. Top-10 candidates ranked by the composite score. Creep life is measured; the remaining properties are five-fold ensemble predictions.", 6.9)
    add_table(doc, tables["comps_top"], "Table 9. Top-10 candidate compositions in at%.", 7.2)
    r1 = top.iloc[0]
    add_paragraph(
        doc,
        f"The highest-ranked candidate combines Ni {r1['Ni']:.2f}, Co {r1['Co']:.2f}, Al {r1['Al']:.2f}, W {r1['W']:.2f}, Ta {r1['Ta']:.2f}, Mo {r1['Mo']:.2f}, Re {r1['Re']:.2f}, Cr {r1['Cr']:.2f}, Ru {r1['Ru']:.2f}, and Hf {r1['Hf']:.2f} at%. Its measured creep life is {r1['creep_real']:.0f} h at {r1['test_temp']:.0f} degC/{r1['test_stress']:.0f} MPa. The predicted solvus is {r1['solvus']:.1f} degC, the processing window is {r1['processing_window']:.1f} degC, the density is {r1['density']:.2f} g/cm3, the harmful phase probability is {r1['phase_class']:.3f}, the gamma-prime size is {r1['size']:.0f}, and the freezing range is {r1['freezing_range']:.1f} degC.",
    )
    add_paragraph(
        doc,
        "The ten recommendations share several design characteristics. They are Ni-rich rather than genuinely Co-base, with Ni mostly between 60 and 65 at% except for the ninth candidate. Re is present in all candidates and is especially high in the highest-creep entries. Ta and W are also consistently high, which is compatible with strengthening and solvus elevation. This ranking illustrates the intended behavior of the composite score: it does not simply maximize creep or solvus, but rewards balanced satisfaction of all engineering criteria.",
    )

    add_heading(doc, "3.6. Robustness, feasibility, and model limitations", 2)
    add_paragraph(
        doc,
        "The workflow contains several safeguards against overconfident model-driven design. All regression targets are normalized by task, SmoothL1 loss and uncertainty weighting reduce sensitivity to outliers and task imbalance, five-fold ensembles are used for inference, and task grouping is justified by feature-weight and gradient-affinity evidence. Most importantly, the inverse-design recommendation keeps measured creep life as an anchor because creep prediction is the most difficult downstream task.",
    )
    add_paragraph(
        doc,
        "The main limitation is that the final recommendations are constrained by the distribution of available high-creep seeds. This improves credibility but reduces novelty relative to a free optimizer. The current phase-class model predicts a probability of harmful phase formation rather than an explicit phase fraction or phase identity. Future work should combine the present model with CALPHAD validation, uncertainty quantification, and targeted experiments.",
    )

    add_heading(doc, "4. Conclusions", 1)
    add_paragraph(
        doc,
        "This work developed an end-to-end superalloy design pipeline that links literature-scale semantic pretraining, multi-task property prediction, data-driven task grouping, and conservative inverse design. The main conclusions are as follows.",
    )
    add_bullet(doc, "A domain-specific tokenizer and PA-MLM pretraining strategy were built from 26,980 superalloy documents and 4.75 million valid tokenized sentences.")
    add_bullet(doc, "The E_pa full multi-task model outperformed the plain E_base model on all seven tasks, reaching R2 = 0.806 for density, 0.736 for creep life, 0.895 for liquidus, 0.946 for gamma-prime size, 0.853 for solidus, 0.875 for solvus, and F1 = 0.930 for harmful phase classification.")
    add_bullet(doc, "Task-gradient analysis showed that E_pa reduces strong inter-task conflicts, whereas E_base relies on a physically coherent temperature-task cluster.")
    add_bullet(doc, "The most difficult target is creep life, motivating the final seed-based screening strategy that uses measured creep life as the reliability anchor.")
    add_bullet(doc, "Seed-based screening produced 10 recommended high-temperature alloys satisfying solvus, density, processing-window, phase, size, and freezing-range constraints.")

    add_heading(doc, "Data availability", 1)
    add_paragraph(
        doc,
        "The processed task datasets, pretrained embeddings, model outputs, figures, and candidate-screening tables are available in the local project directory. The main supervised dataset is data/master_table.csv. Model comparison results are stored in output/results/full_comparison_table.csv and output/results/best_model_per_task_embedding.csv. The final recommendation table is output/results/forward_screen_seedtop10_seeds_T1080-1120_S120-160_life270.csv.",
    )

    add_heading(doc, "Code availability", 1)
    add_paragraph(
        doc,
        "The implementation is organized by workflow phase. Phase 1 scripts build the corpus, tokenizer, and master table. Phase 2 scripts train E_base and E_pa embeddings. Phase 3 scripts train and evaluate multi-task, grouped, single-task, and classical machine-learning models. Phase 4 scripts perform inverse design, forward screening, seed-based screening, and candidate visualization. The key files are code/phase1_data/tokenizer.py, code/phase2_pretrain/model_pa_mlm.py, code/phase3_multitask/model_multitask.py, code/phase3_multitask/task_grouping.py, code/phase4_inverse/forward_screen.py, and code/phase4_inverse/seed_screen.py.",
    )

    add_heading(doc, "Author contributions", 1)
    add_paragraph(doc, "To be completed. Suggested structure: data curation; model development; validation; interpretation; manuscript writing; supervision.")
    add_heading(doc, "Declaration of competing interest", 1)
    add_paragraph(doc, "The authors declare no competing interests. This statement should be confirmed before submission.")
    add_heading(doc, "Acknowledgments", 1)
    add_paragraph(doc, "Funding and computational-resource acknowledgments should be completed by the authors.")
    add_heading(doc, "References", 1)
    add_paragraph(
        doc,
        "Reference list to be completed. Suggested citation groups include data-driven materials design; text mining and language models in materials science; multi-task learning and uncertainty weighting; gradient-based task affinity or task grouping; superalloy creep and gamma-prime strengthening; CALPHAD and phase-stability validation.",
    )

    return doc


def main():
    mt, comp, best, top, corpus_stats, grp_pa, grp_base = load_data()
    tables = make_tables(mt, comp, best, top, corpus_stats, grp_pa, grp_base)
    doc = build_doc(tables, top)
    try:
        doc.save(OUT_PATH)
        saved_path = OUT_PATH
    except PermissionError:
        saved_path = OUT_DIR / f"{OUT_PATH.stem}_updated.docx"
        doc.save(saved_path)
    print(saved_path)


if __name__ == "__main__":
    main()
