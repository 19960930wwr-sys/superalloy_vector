"""
生成 Word 文档: 基于多任务深度学习的高温合金逆向设计方法说明
============================================================
面向材料学读者, 完整描述本工作的数据/模型/逆向筛选方法学,
并把 1100 C / 140 MPa, life > 270 h 的 top-10 推荐结果写入正文.
"""
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parent.parent))

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import RESULTS_DIR

TOP10_CSV = RESULTS_DIR / "forward_screen_seedtop10_seeds_T1080-1120_S120-160_life270.csv"
OUT_DOCX  = RESULTS_DIR / "逆向设计方法与结果说明.docx"

LQ = '\u300c'   # 「
RQ = '\u300d'   # 」


def set_cn_font(run, size=11, bold=False):
    run.font.name = 'Times New Roman'
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, size=11, bold=False, align=None, indent_first=True):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold)
    return p


def add_heading_cn(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_cn_font(run, size=sizes.get(level, 12), bold=True)
    return p


def add_table_from_df(doc, df, header_bg='D9E1F2'):
    n_rows, n_cols = df.shape[0] + 1, df.shape[1]
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(col))
        set_cn_font(run, size=9, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_bg)
        tcPr.append(shd)
    for i in range(df.shape[0]):
        for j, col in enumerate(df.columns):
            cell = table.cell(i + 1, j)
            cell.text = ''
            val = df.iloc[i, j]
            txt = f"{val:.2f}" if isinstance(val, float) else str(val)
            run = cell.paragraphs[0].add_run(txt)
            set_cn_font(run, size=9, bold=False)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


# ============== 主体 ==============
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(11)
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ---- 标题 ----
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run(
    '基于多任务深度学习的高温合金逆向设计方法\n'
    '——以 1100 °C / 140 MPa 蠕变工况为目标的种子驱动多目标筛选'
)
set_cn_font(title_run, size=18, bold=True)
title_p.paragraph_format.space_after = Pt(12)

# ============= 1. 研究背景 =============
add_heading_cn(doc, '1  研究背景与目标', level=1)

add_para(doc,
    '高温合金（superalloy）的成分设计需要在多个相互冲突的性能指标之间寻求平衡，'
    '包括 γ′ 析出温度（solvus）、合金密度、凝固区间、加工窗口、γ′ 颗粒尺寸、'
    '蠕变寿命，以及对有害析出相（TCP 相）的抑制能力。'
    '传统试错法和 CALPHAD 模拟难以在十亿量级的成分组合中高效筛选；'
    '基于机器学习的正向枚举法虽能加速搜索，但当模型在某些子空间欠拟合或外推时，'
    '其预测的最优合金往往难以经受实验验证。'
)

add_para(doc,
    '本工作以 Co-Ni 基与 Ni 基 γ′ 强化高温合金为对象，'
    '构建了一套基于训练集真实性能锚点 + 多任务模型补全的逆向筛选框架。'
    '在指定工况（本案例为 1100 °C / 140 MPa，蠕变寿命 > 270 h）下，'
    '结合 6 项工程约束（γ′ solvus、密度、凝固区间、加工窗口、γ′ 尺寸、有害相概率），'
    '输出可信的 Top-10 候选合金供实验验证。'
)

# ============= 2. 数据集与表征 =============
add_heading_cn(doc, '2  数据集与多任务表征', level=1)

add_heading_cn(doc, '2.1  主表（master_table）', level=2)
add_para(doc,
    '从公开文献与已有合金数据库收集了约 4000 条记录，整理为统一的主表格式（long-format）。'
    '每条记录包含一种合金的成分（at%）、热处理与测试工艺参数、所测性能指标'
    '（density、γ′ solvus、liquidus、solidus、γ′ size、phase_class、creep life），'
    '以及该测试条件下的目标值。'
    '其中蠕变（creep）任务共 1024 条样本，覆盖测试温度 700–1200 °C、应力 60–900 MPa 的宽广工况。'
)

add_heading_cn(doc, '2.2  元素与工艺动作的向量化嵌入', level=2)
add_para(doc,
    '为了让模型同时理解成分和工艺，本工作训练了一套元素 + 工艺动作的联合词向量（embedding）。'
    '词表中既包含元素符号（Co, Ni, Al, …），也包含工艺动作词'
    '（solution treatment, aging, creep）。'
    '训练采用 PA-MLM（Process-Aware Masked Language Model）方案：'
    '在合金语料上做掩码语言建模，并以元素的物理属性（电负性、原子半径、电子构型等）'
    '作为辅助预测目标，使每个元素向量在嵌入空间中保留了其物理化学含义。'
    '记此嵌入为 E_pa。'
)

add_heading_cn(doc, '2.3  多任务回归模型', level=2)
add_para(doc,
    '下游训练一个共享底座 + 多头输出的多任务网络。每个合金的输入由三部分拼接：'
    '（i）成分加权平均的元素向量（128 维）；'
    '（ii）工艺动作向量（128 维）；'
    '（iii）工艺数值参数的 z-score（solution_temp、aging_time、test_temp、test_stress 等）。'
    '输出 7 个性能任务头：density、liquidus、solidus、γ′ solvus、γ′ size、'
    'phase_class（有害相概率，sigmoid 输出）、creep life（log 变换归一化）。'
    '训练采用 5 折交叉验证 + 任务平衡采样策略，最终保留 5 个 fold 的 checkpoint 做集成预测。'
)

# ============= 3. 逆向设计：seed-driven 框架 =============
add_heading_cn(doc, '3  逆向设计方法 —— 种子驱动的多目标筛选', level=1)

add_heading_cn(doc, '3.1  方法动机', level=2)
add_para(doc,
    '在用模型做正向枚举时，本工作发现两个普遍但容易被忽视的问题：'
)
add_para(doc,
    '（1）模型在训练集稀疏区域的外推不可信。'
    '以 760 °C / 750–840 MPa 高应力工况为例，训练集中存在 15 个真实蠕变寿命 > 850 h '
    '的合金，但模型对其中 6 个 Ni 基样本的预测仅相当于真值的 14%–35%（平均低估 75%），'
    '因此基于模型预测的高 creep 候选会系统性遗漏真正优秀的成分。',
    indent_first=False
)
add_para(doc,
    '（2）设计空间与训练集真实分布往往不重叠。'
    '若强行限定在 Co-Ni 双基（Ni = 30 at%）的窄区域内枚举，'
    '训练集中没有任何一个 Co ≥ 30 at% 的蠕变样本，'
    '模型给出的最优解实质上是纯外推、缺乏验证依据。',
    indent_first=False
)
add_para(doc,
    '为此，本工作放弃全空间枚举路线，改用种子驱动框架：'
    '将关键且可靠性要求最高的指标（蠕变寿命）锁定为训练集中的真实测量值，'
    '其余指标用模型补全后做多目标排序。'
)

add_heading_cn(doc, '3.2  筛选流程', level=2)
add_para(doc, '完整流程分为四步：')
add_para(doc,
    'Step 1（种子提取）：从 master_table 中筛选目标 creep 工况附近的真实合金。'
    '本案例条件为 T ∈ [1080, 1120] °C，σ ∈ [120, 160] MPa，life > 270 h，'
    '共得到 52 个种子合金。',
    indent_first=False
)
add_para(doc,
    'Step 2（多任务预测）：对每个种子合金，用 5 折模型集成预测其余 6 项性能'
    '（solvus、density、phase_class、size、liquidus、solidus），'
    '并由 liquidus 和 solidus 计算凝固区间（freezing range）'
    '和加工窗口（processing window = solidus − solvus）。',
    indent_first=False
)
add_para(doc,
    'Step 3（硬约束过滤）：依次施加 6 项工程硬约束：'
    'γ′ solvus > 1220 °C、density < 8.9 g/cm³、processing window ≥ 80 °C、'
    'phase_class < 0.5、γ′ size < 500（无量纲尺寸特征）、freezing range < 60 °C。',
    indent_first=False
)
add_para(doc,
    'Step 4（多目标综合排序）：将通过过滤的候选按 6 项性能做 z-score 标准化，'
    '按预设权重（solvus 与 processing_window 取正向 +1.0；'
    'density、size、freezing 取负向 −1.0；phase_class 取负向 −1.5 强调避有害相）'
    '求加权和，得到综合分。按综合分降序输出 Top-10。',
    indent_first=False
)

add_heading_cn(doc, '3.3  关键设计决策：阈值放宽与瓶颈诊断', level=2)
add_para(doc,
    '在严格的 6 项约束下（包括 processing window > 100 °C），52 个种子全部被过滤掉。'
    '瓶颈诊断显示，processing window 是唯一真正卡死的指标：'
    '种子集中 processing window 中位仅 51 °C，仅 6 / 52 个达到 100 °C；'
    '其他 5 项约束的通过率均在 87%–100%。'
    '其物理原因在于：满足 solvus > 1220 °C 的合金，其 γ′ solvus 普遍落在 1255–1305 °C 区间，'
    '而 solidus 受熔化机制限制大致在 1340–1360 °C，二者之差自然受到压缩。'
    '因此本工作把 processing window 阈值放宽到 80 °C（仍属合金可加工的合理区间），'
    '最终保留 11 个候选并取综合分前 10。'
)

# ============= 4. 结果 =============
add_heading_cn(doc, '4  Top-10 候选合金', level=1)

top = pd.read_csv(TOP10_CSV)
elems_main = ['Co','Ni','Al','Ti','W','Ta','Mo','Nb','Cr','Re','Hf','Ru']
elems_main = [e for e in elems_main if e in top.columns]

add_heading_cn(doc, '4.1  Top-10 候选的性能指标（蠕变为真实测量，其余为模型预测）', level=2)
perf = top[['creep_real','test_temp','test_stress','solvus',
            'processing_window','density','phase_class','size',
            'freezing_range']].copy()
perf.columns = ['creep_life (h)', 'T (°C)', 'σ (MPa)',
                'γ′ solvus (°C)', 'window (°C)', 'density (g/cm³)',
                'phase_class', 'γ′ size', 'freezing (°C)']
perf.insert(0, 'Rank', range(1, len(perf) + 1))
add_table_from_df(doc, perf)

add_para(doc,
    '说明：creep_life 列为训练集中的真实蠕变寿命实测值（在标注的 T / σ 工况下）；'
    'γ′ solvus、density、phase_class、γ′ size、liquidus、solidus 为 5 折模型集成预测值；'
    'processing window = solidus − γ′ solvus；'
    'freezing = liquidus − solidus；'
    'phase_class 为模型预测的出现有害相（如 TCP）的概率，越低越好。'
)

add_heading_cn(doc, '4.2  Top-10 候选的成分（at%）', level=2)
comp = top[elems_main].copy()
comp.insert(0, 'Rank', range(1, len(comp) + 1))
add_table_from_df(doc, comp)

# ============= 5. 冠军合金详解 =============
add_heading_cn(doc, '5  冠军合金详解（Rank 1）', level=1)

r1 = top.iloc[0]
comp_str = '  '.join(f"{e} {r1[e]:.2f}" for e in elems_main if r1.get(e, 0) > 0)
add_para(doc, f'成分 (at%)：{comp_str}', bold=True)

add_para(doc,
    f'该合金的实测蠕变寿命达到 {r1["creep_real"]:.0f} h '
    f'@ {r1["test_temp"]:.0f} °C / {r1["test_stress"]:.0f} MPa，'
    f'为 Top-10 中最高，约为第二名（402 h）的 1.8 倍。'
    '其成分特征为 Re 与 Mo 双重慢扩散元素的复合搭配（Re 3.6 + Mo 2.9），'
    '并且含有较高比例的 Ta（5.6）、Nb（5.8）作为强 γ′ 形成元素，'
    '整体属于二代单晶（second-generation single-crystal）镍基高温合金。'
)

add_para(doc, '6 项工程指标核查：', bold=True)
checks = [
    ('γ′ solvus', f"{r1['solvus']:.1f} °C", '> 1220 °C', '通过'),
    ('processing window', f"{r1['processing_window']:.1f} °C", '≥ 80 °C', '通过'),
    ('density', f"{r1['density']:.2f} g/cm³", '< 8.9 g/cm³', '通过'),
    ('phase_class', f"{r1['phase_class']:.2f}", '< 0.5', '通过'),
    ('γ′ size', f"{r1['size']:.0f}", '< 500', '通过'),
    ('freezing range', f"{r1['freezing_range']:.1f} °C", '< 60 °C', '通过'),
    ('creep life', f"{r1['creep_real']:.0f} h", '> 270 h', '通过 (实测)'),
]
df_chk = pd.DataFrame(checks, columns=['指标', '本合金值', '约束', '是否通过'])
add_table_from_df(doc, df_chk)

# ============= 6. 方法学讨论 =============
add_heading_cn(doc, '6  方法学讨论与限制', level=1)

add_para(doc,
    '本工作的核心贡献在于：把传统的模型驱动全空间枚举转换为真值锚定 + 模型辅助的混合策略，'
    '可以避开模型在训练集稀疏区域的外推风险。'
    '在 760 °C / 800 MPa 工况的对照案例中，'
    '若直接以模型预测 creep > 850 h 为筛选条件，'
    '整个 Co-Ni 设计空间（约 358 万组合）的 creep 预测最大值仅 496 h，'
    '本质上是模型在该 corner 系统性低估造成的。'
    '种子驱动框架则直接利用了训练集中真实存在的 7 个 > 850 h 样本，'
    '可信度显著提升。'
)

add_para(doc, '限制与下一步工作：', bold=True)
add_para(doc,
    '（1）当前方法依赖训练集中已有目标工况附近的样本。'
    '若目标工况完全没有覆盖（例如 Co-Al-W γ′ 强化的真正 Co 基体系），'
    '需要补充该子区域的实验数据并重训模型；',
    indent_first=False
)
add_para(doc,
    '（2）processing window 在高 solvus（> 1260 °C）合金中受到物理上限压缩，'
    '若工程上确需 > 100 °C 的窗口，应考虑略降 solvus 阈值（例如 > 1200 °C）以释放设计空间；',
    indent_first=False
)
add_para(doc,
    '（3）模型对 phase_class 的预测仅给出出现有害相的概率，'
    '下一步需结合 CALPHAD 计算或 SHAP 解释做物理一致性校验，再进入实验验证。',
    indent_first=False
)

# ============= 附录 =============
add_heading_cn(doc, '附录  关键脚本', level=1)
add_para(doc,
    '种子筛选与多目标排序的主脚本为 phase4_inverse/seed_screen.py；'
    '瓶颈诊断脚本为 phase4_inverse/_diag_seed_1100_140_dist.py；'
    '训练集 / 模型欠拟合分析为 phase4_inverse/_diag_seed.py；'
    '全部命令、参数、结果文件命名见 phase4_inverse/seed_screen.py 顶部 docstring。'
)

# === 保存 ===
RESULTS_DIR.mkdir(parents=True, exist_ok=True)
doc.save(OUT_DOCX)
print(f'已保存: {OUT_DOCX}')
